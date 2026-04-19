"""
DOCX Writer — generates a professional Word document from BRD & WBS data.
"""

import io
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------

_BRAND_DARK = RGBColor(0x0F, 0x17, 0x2A)   # deep navy
_BRAND_ACCENT = RGBColor(0x38, 0x7A, 0xFF)  # electric blue
_BRAND_LIGHT = RGBColor(0xF0, 0xF4, 0xFF)   # light blue tint
_TABLE_HEADER_BG = "1E3A5F"                  # dark blue for table headers
_TABLE_ALT_ROW = "EAF0FB"                    # alternating row tint
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def create_docx(
    project_brief: str,
    brd_text: str,
    wbs_df: pd.DataFrame,
    agile_df: pd.DataFrame,
    risk_df: pd.DataFrame,
    cost_df: pd.DataFrame,
    methodology: str,
) -> io.BytesIO:
    """
    Build a professional DOCX report and return it as an in-memory buffer.

    Parameters
    ----------
    project_brief : str
        The original user-provided project description.
    brd_text : str
        The raw BRD text (with markdown section headers).
    wbs_df : pd.DataFrame
        Already-parsed WBS DataFrame with columns
        [Phase, Task, Estimated Days].

    Returns
    -------
    io.BytesIO
        A seeked-to-zero buffer containing the .docx bytes.
    """
    doc = Document()

    # -- Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    _add_title_page(doc, project_brief)
    _add_brd_section(doc, brd_text)
    if methodology and methodology.startswith("Waterfall"):
        _add_wbs_section(doc, wbs_df)
    else:
        _add_agile_section(doc, agile_df)
    _add_risk_section(doc, risk_df)
    _add_cost_section(doc, cost_df)
    _add_footer(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _add_title_page(doc: Document, project_brief: str) -> None:
    """Create a styled title page."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph("")

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("APEX Lite")
    run.font.size = Pt(36)
    run.font.color.rgb = _BRAND_ACCENT
    run.bold = True

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI-Powered Project Planner")
    run.font.size = Pt(16)
    run.font.color.rgb = _BRAND_DARK
    run.italic = True

    # Divider line
    doc.add_paragraph("─" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project brief
    brief_heading = doc.add_paragraph()
    brief_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brief_heading.add_run("Project Brief")
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = _BRAND_DARK

    brief_body = doc.add_paragraph()
    brief_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brief_body.add_run(project_brief)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()


def _add_brd_section(doc: Document, brd_text: str) -> None:
    """Parse the BRD markdown into proper Word headings and bullets."""
    _section_heading(doc, "Business Requirements Document")

    for line in brd_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        # Markdown heading → Word heading
        if stripped.startswith("## "):
            h = doc.add_heading(stripped.lstrip("# ").strip(), level=2)
            for run in h.runs:
                run.font.color.rgb = _BRAND_DARK
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped.lstrip("# ").strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = _BRAND_DARK
        elif stripped.startswith("- ") or stripped.startswith("• "):
            bullet_text = stripped.lstrip("-•").strip()
            p = doc.add_paragraph(bullet_text, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(11)
        else:
            p = doc.add_paragraph(stripped)
            for run in p.runs:
                run.font.size = Pt(11)

    doc.add_page_break()


def _add_wbs_section(doc: Document, wbs_df: pd.DataFrame) -> None:
    """Add the WBS as a professionally formatted table."""
    _section_heading(doc, "Work Breakdown Structure")

    if wbs_df.empty:
        doc.add_paragraph("WBS data could not be parsed into a table.")
        return

    cols = list(wbs_df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    for idx, col_name in enumerate(cols):
        header_cells[idx].text = col_name
        _style_cell(header_cells[idx], bold=True, bg=_TABLE_HEADER_BG, fg=_WHITE, size=11)

    # Data rows
    for row_idx, (_, row) in enumerate(wbs_df.iterrows()):
        cells = table.add_row().cells
        bg = _TABLE_ALT_ROW if row_idx % 2 == 0 else None
        for col_idx, col_name in enumerate(cols):
            cells[col_idx].text = str(row[col_name])
            _style_cell(cells[col_idx], size=10, bg=bg)

    # Summary row
    total_days = _sum_days(wbs_df)
    if total_days > 0:
        summary_cells = table.add_row().cells
        summary_cells[0].text = ""
        summary_cells[1].text = "TOTAL"
        summary_cells[2].text = f"{total_days} days"
        for c in summary_cells:
            _style_cell(c, bold=True, size=11)
            
    doc.add_page_break()

def _add_agile_section(doc: Document, agile_df: pd.DataFrame) -> None:
    """Add the Agile Backlog as a professionally formatted table."""
    _section_heading(doc, "Product Backlog (Agile)")

    if agile_df is None or agile_df.empty:
        doc.add_paragraph("Agile Backlog data could not be parsed into a table.")
        doc.add_page_break()
        return

    cols = list(agile_df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    for idx, col_name in enumerate(cols):
        header_cells[idx].text = col_name
        _style_cell(header_cells[idx], bold=True, bg=_TABLE_HEADER_BG, fg=_WHITE, size=11)

    # Data rows
    for row_idx, (_, row) in enumerate(agile_df.iterrows()):
        cells = table.add_row().cells
        bg = _TABLE_ALT_ROW if row_idx % 2 == 0 else None
        for col_idx, col_name in enumerate(cols):
            cells[col_idx].text = str(row[col_name])
            _style_cell(cells[col_idx], size=10, bg=bg)
            
    # Summary row
    total_points = 0
    if "Story Points" in agile_df.columns:
        for v in agile_df["Story Points"]:
            try:
                total_points += int(v)
            except (ValueError, TypeError):
                pass
                
    if total_points > 0:
        summary_cells = table.add_row().cells
        summary_cells[0].text = ""
        summary_cells[1].text = "TOTAL POINTS"
        summary_cells[2].text = str(total_points)
        if len(cols) > 3:
            summary_cells[3].text = ""
        for idx in range(len(cols)):
            _style_cell(summary_cells[idx], bold=True, size=11)

    doc.add_page_break()

def _add_risk_section(doc: Document, risk_df: pd.DataFrame) -> None:
    """Add the Risk Register as a professionally formatted table."""
    _section_heading(doc, "Risk Register")

    if risk_df is None or risk_df.empty:
        doc.add_paragraph("Risk data could not be parsed into a table.")
        doc.add_page_break()
        return

    cols = list(risk_df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    for idx, col_name in enumerate(cols):
        header_cells[idx].text = col_name
        _style_cell(header_cells[idx], bold=True, bg=_TABLE_HEADER_BG, fg=_WHITE, size=11)

    # Data rows
    for row_idx, (_, row) in enumerate(risk_df.iterrows()):
        cells = table.add_row().cells
        bg = _TABLE_ALT_ROW if row_idx % 2 == 0 else None
        for col_idx, col_name in enumerate(cols):
            cells[col_idx].text = str(row[col_name])
            _style_cell(cells[col_idx], size=10, bg=bg)

    doc.add_page_break()

def _add_cost_section(doc: Document, cost_df: pd.DataFrame) -> None:
    """Add the Cost Estimation as a professionally formatted table."""
    _section_heading(doc, "Cost Estimation")

    if cost_df is None or cost_df.empty:
        doc.add_paragraph("Cost data could not be parsed into a table.")
        doc.add_page_break()
        return

    cols = list(cost_df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    for idx, col_name in enumerate(cols):
        header_cells[idx].text = col_name
        _style_cell(header_cells[idx], bold=True, bg=_TABLE_HEADER_BG, fg=_WHITE, size=11)

    # Data rows
    for row_idx, (_, row) in enumerate(cost_df.iterrows()):
        cells = table.add_row().cells
        bg = _TABLE_ALT_ROW if row_idx % 2 == 0 else None
        for col_idx, col_name in enumerate(cols):
            cells[col_idx].text = str(row[col_name])
            _style_cell(cells[col_idx], size=10, bg=bg)
            
    doc.add_page_break()


def _add_footer(doc: Document) -> None:
    """Add a subtle footer note."""
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by APEX Lite — AI Project Planner")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True


# ---------------------------------------------------------------------------
# Micro-helpers
# ---------------------------------------------------------------------------

def _section_heading(doc: Document, text: str) -> None:
    """Add a styled top-level section heading."""
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = _BRAND_ACCENT
        run.font.size = Pt(20)


def _style_cell(
    cell,
    bold: bool = False,
    bg: str | None = None,
    fg: RGBColor | None = None,
    size: int = 10,
) -> None:
    """Apply font and shading to a table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.bold = bold
            if fg:
                run.font.color.rgb = fg

    if bg:
        shading = cell._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(
            qn("w:shd"),
            {"val": "clear", "color": "auto", "fill": bg},
        )
        shading.append(shading_elem)


def _sum_days(df: pd.DataFrame) -> int:
    """Sum 'Estimated Days' column, ignoring non-numeric values."""
    total = 0
    for val in df["Estimated Days"]:
        try:
            total += int(val)
        except (ValueError, TypeError):
            pass
    return total